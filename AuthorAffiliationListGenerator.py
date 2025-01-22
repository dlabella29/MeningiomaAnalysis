import pandas as pd
from docx import Document

# Function to add a run with superscript text
def add_superscript(paragraph, text):
    run = paragraph.add_run(text)
    run.font.superscript = True

# Function to add text (including commas and spaces as needed)
def add_text(paragraph, text):
    paragraph.add_run(text)

# Function to save the author list to a Word document, with ordered superscripts
def save_author_list_to_word(file_path, output_doc_path):
    # Load the CSV file
    df = pd.read_csv(file_path)

    # Check if the required columns exist
    required_columns = ['First Name', 'Last Name', 'Middle Name(s)', 'Affiliation1', 'Affiliation2', 'Affiliation3', 'Affiliation4']
    for column in required_columns:
        if column not in df.columns:
            raise ValueError(f"Column '{column}' is missing in the input file.")

    # Create a new Document
    doc = Document()
    doc.add_heading('Author List and Affiliations', 0)

    # Dictionary to hold affiliations and their IDs
    affiliations = {}
    affiliation_counter = 1

    # Start the author list in the same paragraph
    paragraph = doc.add_paragraph(style='Normal')

    for index, row in df.iterrows():
        # Construct the author's name
        first_name = row['First Name']
        last_name = row['Last Name']
        middle_name = row['Middle Name(s)'] if pd.notna(row['Middle Name(s)']) else ''
        middle_initial = f" {middle_name[0]}." if middle_name else ''
        name = f"{first_name}{middle_initial} {last_name}"

        # Only add a comma and space before the name if it's not the first author
        if index > 0:
            add_text(paragraph, ', ')

        # Add the author's name
        add_text(paragraph, name)

        # List to track an author's affiliations for superscript, sorted
        author_affiliations = []

        # Process each possible affiliation column
        for aff_col in ['Affiliation1', 'Affiliation2', 'Affiliation3', 'Affiliation4']:
            affiliation = row[aff_col]
            if pd.notna(affiliation):
                if affiliation not in affiliations:
                    affiliations[affiliation] = affiliation_counter
                    affiliation_counter += 1
                author_affiliations.append(affiliations[affiliation])

        # Sort the affiliation numbers for each author
        author_affiliations.sort()

        # Add the affiliations as superscripts with commas between them
        for i, aff_num in enumerate(author_affiliations):
            if i > 0:
                add_superscript(paragraph, ',')
            add_superscript(paragraph, str(aff_num))

    # Add affiliations at the end
    doc.add_page_break()
    doc.add_heading('Affiliations', level=1)
    for affiliation, num in sorted(affiliations.items(), key=lambda x: x[1]):
        doc.add_paragraph(f"{num}. {affiliation}")

    # Save the document
    doc.save(output_doc_path)

# Specify the path to your input CSV file and the output .docx file
file_path = r"C:\Users\dlabe\Downloads\Meningioma authors BraTS-MEN-RT_SciData.csv"
output_doc_path = r"C:\Users\dlabe\Downloads\Author_List_and_Affiliations BraTS-MEN-RT_SciData.docx"
# Generate the updated Word document
save_author_list_to_word(file_path, output_doc_path)

print("Author list and affiliations document created successfully!")

